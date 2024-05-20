import os
import base64
import re
from docx import Document
from PIL import Image
from io import BytesIO
import matplotlib.pyplot as plt
import matplotlib

# Matplotlibning TeX-ni ishlatishiga ishonch hosil qiling
matplotlib.rc('text', usetex=True)


def read_docx_and_convert_to_html(docx_path):
    document = Document(docx_path)
    html_content = "<html><body>"

    for para in document.paragraphs:
        html_content += f"<p>{para.text}</p>"

    for table in document.tables:
        html_content += "<table border='1'>"
        for row in table.rows:
            html_content += "<tr>"
            for cell in row.cells:
                html_content += f"<td>{cell.text}</td>"
            html_content += "</tr>"
        html_content += "</table>"

    html_content += "</body></html>"
    return html_content


def convert_formula_to_image(formula):
    fig = plt.figure(figsize=(2, 2))
    text = f"${formula}$"
    plt.text(0.5, 0.5, text, fontsize=20, ha='center', va='center')
    plt.axis('off')

    buf = BytesIO()
    plt.savefig(buf, format='png', bbox_inches='tight', pad_inches=0.1)
    plt.close(fig)
    buf.seek(0)
    img_str = base64.b64encode(buf.getvalue()).decode('utf-8')
    return img_str


def find_and_replace_formulas(html_content):
    # Regex bilan formulalarni topish
    formula_pattern = re.compile(r'\$\$(.*?)\$\$')
    formulas = re.findall(formula_pattern, html_content)

    for formula in formulas:
        img_str = convert_formula_to_image(formula)
        img_tag = f'<img src="data:image/png;base64,{img_str}" />'
        html_content = html_content.replace(f"$${formula}$$", img_tag)

    return html_content


def main(docx_path, output_txt_path):
    html_content = read_docx_and_convert_to_html(docx_path)
    html_content_with_images = find_and_replace_formulas(html_content)

    # Natijani faylga yozish
    with open(output_txt_path, 'w') as f:
        f.write(html_content_with_images)


docx_path = '/Users/auzcoder/Desktop/testuchun/pythonProject/Moodle-quiz-create/math.docx'  # DOCX fayl manzili
output_txt_path = '/Users/auzcoder/Desktop/testuchun/pythonProject/Moodle-quiz-create/output.txt'  # TXT fayl manzili

main(docx_path, output_txt_path)
